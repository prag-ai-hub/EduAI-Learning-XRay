"""Build an evidence-led acceptance register; never equate untested with pass.
Preset: standard_business_brief; masthead: memo_masthead (no rule).
Named overrides: Title 23pt/after4, Subtitle 12pt/after10;
case labels 9pt gray; metadata 10pt/after4. No tables or fake lists.
"""
from pathlib import Path
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT=Path(__file__).parent
# Area, test, procedure, expected, actual, status. Evidence is limited to the
# stated scenario; historical tests are labelled rather than silently retested.
cases=[
('Authentication','Existing teacher session','Reload the published HPC /app page.','Existing teacher reaches workspace without a blank page.','Workspace restored; no fresh credentials were entered.','PASS'),
('Authentication','Fresh email/password login','Use a dedicated test account; log in, reload, log out, and retry an invalid password.','Valid login succeeds; invalid login fails clearly; logout removes access.','Fresh credential flow not executed. Existing-session restoration is not equivalent.','NOT TESTED'),
('Authentication','Email verification and recovery','Register a dedicated account; use a new confirmation email and password-recovery link.','Links return to HPC; expired/reused links show a safe error.','No fresh inbox-controlled account available in this execution.','BLOCKED'),
('Authentication','Google OAuth','Complete new-user and returning-user consent and callback.','Correct HPC account/session created; reload and logout work.','Provider disabled and invalid Client ID previously confirmed; user deferred Google work.','DEFERRED'),
('Authentication','Microsoft OAuth','Complete Microsoft consent and HPC callback.','User reaches the HPC workspace.','Provider previously confirmed disabled; no Microsoft OAuth flow executed.','FAIL'),
('Prompt 1','Middle framework library','Open Holistic Progress on published version 36.','Middle sections, domains and performance levels load from the Middle framework.','Sections, eight domains and Beginner/Proficient/Advanced loaded. Earlier wrong-framework display corrected.','PASS'),
('Prompt 1','New Middle learner','Create TEST ONLY QA Middle 20260831, roll QA-M-0831, Grade 7, year 2026-27; reopen HPC.','Profile persists with correct stage.','Creation succeeded; Grade 7 and Middle Stage appeared after reopening.','PASS'),
('Prompt 1','Immediate learner-list refresh','After successful creation, inspect Choose HPC learner without leaving the module.','New learner is immediately selectable in downstream sections.','Existing downstream list omitted the new learner until module was reopened.','FAIL'),
('Prompt 1','Existing student linkage','Choose an existing Academic X-Ray student and add an HPC profile without duplicating the student.','Same student ID and roster context are retained.','Code review found learner POST always creates a new student; no existing-student choice tested.','FAIL'),
('Prompt 1','Goal entry and persistence','For the Grade 7 test learner choose Academic goal, enter a synthetic goal, save and reopen.','Goal persists under the learner.','Academic Goal and exact test text appeared after reopening.','PASS'),
('Prompt 1','Science mapping cascade','Choose Science, SCCG1, SCC1.1; save official mapping for review and reopen.','Dependent options load and mapping remains awaiting review.','Goal and competency options loaded; saved Science mapping appeared as Teacher Review Required. Individual outcome mapping not included in this pass.','PASS'),
('Prompt 1','All stage boundaries','Test Balvatika/0,2,3,5,6,8,9,12 and invalid grades with stage-specific fields.','Correct approved stage and only applicable questions/rules are shown.','Grade 7 and an existing Grade 10 were sampled; full boundary matrix was not run.','NOT TESTED'),
('Prompt 1','Secondary official mapping coverage','Select the Grade 10 test learner and inspect its own complete subject/goal/competency/outcome catalogue.','Secondary-specific approved mappings are usable.','Earlier Secondary mapping lists were empty; complete corrected Secondary coverage has not been proven.','BLOCKED'),
('Prompt 1','Attendance, context and validation','Save attendance 0 and 100, reject -1 and 101; verify low-attendance reason and resources after reopen.','Boundary validation and persistent context.','Existing Grade 10 context rendered; edit/negative boundaries not executed.','NOT TESTED'),
('Prompt 1','Version retention and feature flag','Finalize under one framework; change active version in a disposable tenant; also test HPC disabled.','Old report version unchanged; disabled school retains Academic X-Ray behaviour.','No disposable tenant/framework migration exercised.','NOT TESTED'),
('Prompt 2','Activity and rubric creation','Create an activity with date, term, pedagogy and method; save three rubric descriptors and mappings.','All fields persist and remain correctly linked.','Existing test activities rendered; new full activity/rubric lifecycle not executed in this pass.','NOT TESTED'),
('Prompt 2','Teacher observation and bulk entry','Record level/confidence/attachment for one learner and a class with an individual override.','Attribution, class assignment and draft moderation preserved.','Existing observation shown; full capture and bulk class flow not executed.','NOT TESTED'),
('Prompt 2','Structured reflection','Enter reflection plus learning, practice-needed and help-needed for the synthetic Grade 10 learner; save.','Separate learner reflection enters moderation.','Save succeeded in earlier 31 August browser pass; exact reflection persisted after reopen. Individual subfield read-back not verified.','PASS'),
('Prompt 2','Moderation approval and timeline','Approve only the synthetic reflection; reopen dashboard.','Approved entry appears; unapproved entries stay out.','Timeline showed approved count 1 and exact synthetic reflection in earlier same-day pass.','PASS'),
('Prompt 2','Parent partnership entry','Select Grade 7 test learner; enter synthetic parent feedback and home commitment; save and reopen.','Parent perspective remains separate and awaiting moderation.','Save succeeded; exact parent feedback appeared with Approve/Exclude controls. Commitment read-back not verified.','PASS'),
('Prompt 2','Peer feedback entry','Choose Grade 7 test learner; enter synthetic peer feedback; save and reopen.','Separate peer entry remains moderated.','Save succeeded and exact peer text appeared in moderation. Reviewer was optional in this form.','PASS'),
('Prompt 2','Peer exclusion','Exclude only the synthetic peer entry.','Entry cannot contribute to approved evidence.','UI reported Evidence excluded. Persistence and exclusion from every report still require retest.','PARTIAL'),
('Prompt 2','Evidence sufficiency','Approve an unmapped reflection with no teacher/peer/parent evidence.','One mapping gap and missing perspectives are reported.','Dashboard displayed mapping gaps 1 and missing teacher, peer and parent perspectives in earlier same-day pass.','PASS'),
('Prompt 2','Portfolio upload/download','Upload allowed PDF/image/text files for a test learner; reopen attachment and compare bytes; reject unsafe type/oversize.','Private persisted file; exact download; controlled validation.','Existing portfolio entry rendered, but file bytes and new upload/download not verified.','NOT TESTED'),
('Prompt 2','Evidence mapping and academic reference','Map one evidence item to outcome and ability; reference an Academic X-Ray result without blending scores.','Correct context and provenance retained.','Learner-level mapping tested; evidence-level mapping and academic reference flow not executed.','NOT TESTED'),
('Prompt 2','Contributor sharing lifecycle','Create peer/parent links for synthetic learners; submit as named contributors; expire/revoke; verify history.','Intended child/task only; identity, attribution, expiry and revocation enforced.','Existing history rendered. External contributor and revoked-token scenarios not run.','NOT TESTED'),
('Prompt 3','Middle official threshold in UI','For Grade 7 select Awareness/Self, enter count 3 and save.','Self becomes Proficient; peer/teacher remain unrecorded.','UI displayed self 3/6 proficient; other perspectives remained not recorded.','PASS'),
('Prompt 3','All count boundaries and no averaging','Run scoring unit tests for 0..6, invalid counts, missing and conflicting perspectives.','0..2 Beginner,3..4 Proficient,5..6 Advanced; reject invalid values; no aggregate.','Four new automated cases passed, including NaN, Infinity, fractions, null, strings and separated perspectives.','PASS'),
('Prompt 3','Secondary rule selection','Open Grade 10 progress and inspect displayed rule.','Secondary learner must not use Middle-only six-statement scoring.','Earlier browser test displayed Middle rule for Grade 10; route remains hardcoded to Middle. No fix included in v36.','FAIL'),
('Prompt 3','Progress wheel and evidence detail','Record self/peer/teacher evidence, inspect labels and wheel, open mapped evidence.','All required perspectives accessible without color dependence; evidence opens.','Basic progress controls rendered. Full wheel, click-through and perspective correspondence not verified.','NOT TESTED'),
('Prompt 3','Override audit and sensitive feedback','Apply teacher override with reason; inspect audit. Save restricted barrier and inspect parent view.','Override auditable; restricted note not shared; no clinical inference.','Not executed with audit and parent-view verification.','NOT TESTED'),
('Prompt 3','Interventions and follow-up integration','Create support action from mapped evidence; open Academic X-Ray Interventions and follow-up.','Same linked support workflow, no duplicate engine or academic score mutation.','Interventions navigation rendered; full linked lifecycle not verified.','NOT TESTED'),
('Prompt 4','Record detail rendering','Open existing TEST ONLY local water inquiry for Grade 10.','Edit, context, schedule and barrier controls load.','Record detail, academic term/class, milestone and barrier forms loaded.','PASS'),
('Prompt 4','Milestone persistence','Add TEST ONLY 2026-08-31 compare fictional observations milestone, save and reopen.','Milestone retained on correct inquiry.','Saved and reopened with planned status in earlier same-day pass.','PASS'),
('Prompt 4','Group project lifecycle','Create multi-learner Secondary project, roles, stages1/2/3, review and final rubric.','Correct members and evidence through full completion.','Not executed. Existing inquiry is not a substitute for this case.','NOT TESTED'),
('Prompt 4','Inquiry and interaction completion','Complete individual inquiry and each supported classroom interaction type, including reflection and assessment.','All required fields, hours and evidence retained.','Existing in-progress inquiry rendered only; full completion not tested.','NOT TESTED'),
('Prompt 4','Context, barrier and official score','Edit term/class, mappings, barrier lifecycle and finalized rubric per record.','Context retained; unsupported scores rejected; official rule version captured.','Fields render; complete mutation/validation lifecycle not tested.','NOT TESTED'),
('Prompt 4','Online-course proof','Record hours/completion and certificate; download from record; test invalid file and negative hours.','Valid private proof linked; invalid values rejected.','Not executed.','NOT TESTED'),
('Prompt 4','Stage gating and credits','Attempt younger-stage project creation and invalid membership; check required credit totals.','Younger stages and wrong-stage members denied server-side; valid credit data retained.','Client stage-gate code present, but full positive/negative API tests not run.','NOT TESTED'),
('Prompt 5','Negative readiness checks','Run checks for incomplete Grade 10 test record.','Missing assessment, completed activity, official rubric, mappings and teacher approval flagged.','Checks refreshed and flagged missing requirements in earlier same-day browser pass.','PASS'),
('Prompt 5','Approved-evidence draft','Generate annual narrative from only approved evidence; edit and save as draft.','No unapproved or fabricated evidence; teacher retains authority.','Not executed through full narrative generation/review.','NOT TESTED'),
('Prompt 5','Successful immutable finalization','Complete mandatory stage evidence and official rubric; attest, finalize, then attempt edit/refinalize.','Immutable version and evidence snapshot; safe duplicate request handling.','Representative complete mapped evidence and stage-correct scoring are not ready; no successful finalization proven.','BLOCKED'),
('Prompt 5','Final PDF and parent report','Generate final branded PDF, verify signature/content/version; share final report and test restricted data.','Valid PDF and learner-scoped parent view without restricted notes.','Automated PDF tests pass for existing report utility; live final HPC PDF and sharing not tested.','BLOCKED'),
('Prompt 5','Principal dashboard','Use authorised principal plus teacher account; compare grade/term aggregates and sample sizes.','Aggregate default, correct permissions, no rankings or combined score.','Only teacher session available; principal role workflow not tested.','BLOCKED'),
('Security','Anonymous private API access','Send unsigned GETs to learners,evidence,progress,applied-learning,annual-reports,principal-dashboard,shares.','HTTP401 without private records.','All seven returned401 in this published-v36 run.','PASS'),
('Security','Cross-school read/write isolation','Use two controlled school accounts; swap learner/evidence/report/file IDs for read and write attempts.','Denied without data or changes; RLS independently enforced.','Second test school/account not available; source checks cannot certify RLS.','BLOCKED'),
('Security','Parent/peer token isolation','Try wrong learner, modified, expired, revoked and replayed tokens; inspect file access.','Only authorised intended scope; invalid tokens denied.','Not executed.','NOT TESTED'),
('Security','Role and feature-flag enforcement','Use disabled user, student, parent, teacher and principal; exercise mutation APIs with HPC off.','Each action respects role, assignment and school flag.','Only active teacher tested; multi-role/flag negative tests remain.','NOT TESTED'),
('Responsive and accessibility','Narrow layout and full navigation','At measured viewport443px/client428px open HPC; inspect overflow and navigation.','No document horizontal overflow; all sections reachable.','ScrollWidth428 equals clientWidth428; all11 navigation items present. Screenshot inspected for the sampled form only.','PASS'),
('Responsive and accessibility','Desktop/laptop/tablet/phone matrix','Repeat key forms at1440x930,1366x768,1024x768,768x1024,393x844 and360x800.','Readable layout and working controls at each measured size/orientation.','Previous viewport overrides did not change actual app dimensions. Only measured narrow layout verified; other sizes untested.','BLOCKED'),
('Responsive and accessibility','Keyboard, screen reader and real devices','Test Tab order, labels, focus, errors, zoom200%, touch and Safari/Chrome/Edge on real phone/tablet.','Operable without mouse; no clipped forms or inaccessible actions.','Not executed across real devices or screen readers.','NOT TESTED'),
('Regression and operations','All workspace section navigation','Click Home,Work,Review,X-Ray,Interventions,Students,Resources,Achievements,Reports,Settings; open HPC.','Each renders the intended heading, not blank page.','All11 section headings rendered in published v36. This is smoke coverage, not each inner action.','PASS'),
('Regression and operations','Full Academic X-Ray cycle','Create assessment/blueprint; upload multi-student work; OCR, validate, grade, review, generate4PDFs+ZIP, resources and parentQR.','Existing cycle and credits remain correct in HPC copy.','Account displays0credits; complete AI-backed flow not executed.','BLOCKED'),
('Regression and operations','Weak connectivity and large class','Interrupt save/upload and restore; retry; load representative large class and inspect responsiveness.','No duplicates/data loss; clear recovery; acceptable measured performance.','No fault injection or load test performed.','NOT TESTED'),
('Regression and operations','Automated suite and build','Run node --test tests/*.test.mjs; build exact publication source.','No test/build failures.','72 automated tests passed; successful build reused for v36. Many assertions are source-contract tests, not live journeys.','PASS'),
]

# Retest updates preserve case identifiers and distinguish live checks from doubles.
updates={
'Inquiry and interaction completion':('PARTIAL','Published v40: created completed synthetic Discussion, Debate, Simulation/role play, Laboratory and Digital learning records with hours 0,1,1.5,2,0.5. Found stale edit fields when switching records. Fixed in v41 and browser-verified Digital learning, original inquiry and zero-hour Discussion each load their own fields. No browser errors in sampled retest. Full inquiry completion and final rubric remain unverified.'),
'Microsoft OAuth':('FAIL','Read-only HPC Auth settings check: Microsoft/Azure false, email true. Entra client ID, tenant and secret are unavailable; no provider configuration changed.'),
'Immediate learner-list refresh':('PASS','Published v37: new Grade 8 / 2027-28 profile appeared immediately while previous learner selection remained selected.'),
'Existing student linkage':('PASS','Published v37: reused existing synthetic student, roster remained four entries; same-year duplicate was rejected.'),
'Secondary rule selection':('PASS','Published v37: Grade 10 cannot open the Middle six-count form and receives Secondary rubric guidance. This is stage isolation, not proof of complete Secondary scoring.'),
'Attendance, context and validation':('PASS','Live synthetic learner: 0 and 100 saved; 100 and context/resources survived reopening. HTML constraints rejected -1/101. Direct negative API not tested.'),
'Activity and rubric creation':('PARTIAL','Activity created and persisted. Version 38 browser retest reloaded the saved beginner descriptor for the synthetic activity. Full term/mapping lifecycle remains.'),
'Teacher observation and bulk entry':('PARTIAL','Versions 39 and 45: two synthetic learners retained individual notes and confidence after reopening. Published UI now omits null grades from the class filter, disables ungraded learners, enables Middle performance levels only for Grades 6-8, and disables them for the sampled Grade 10 learner. Stable retry IDs pass route tests. Full-class scale, attachment and live interruption tests remain.'),
'Peer exclusion':('PARTIAL','Excluded entry stayed out of live moderation/approved timeline. Real finalization handler with a database double excluded pending/excluded content from narrative and snapshot. Live final PDF is still skipped/blocked.'),
'Portfolio upload/download':('PARTIAL','Real download handler with storage double preserved exact synthetic bytes and enforced school ownership. Fixed use of current teacher folder instead of recorded uploader. Live new upload and byte round-trip remain.'),
'Online-course proof':('PARTIAL','Real proof/download handlers with doubles preserved certificate bytes, rejected negative hours and retained zero after local fix. Live upload/completion journey remains.'),
'Interventions and follow-up integration':('PARTIAL','Version 39: existing completed HPC support action and review date are visible in Interventions; a planned synthetic action was started there. Same HPC endpoint is used, no academic snapshot copy. Full follow-up-evidence lifecycle remains.'),
'Approved-evidence draft':('PASS','Published v38 generated drafts from approved evidence only. On v39 the synthetic Grade 10 draft was teacher-edited, saved and reopened: narrative, strengths, support next steps and teacher notes all retained their exact QA40 text. Draft remained separate from finalization. This pass covers draft generation/edit/persistence, not immutable final reporting.'),
'Evidence mapping and academic reference':('PARTIAL','Published v39: mapped the synthetic Grade 7 BULK39 observation to Science, SCCG1, SCC1.1 and Awareness; approved only that test observation. After reopening, its evidence-detail view displayed Awareness and SCC1.1. No learning-outcome option was available for this competency; Academic X-Ray reference and no-blending journey remain untested.'),
'Progress wheel and evidence detail':('PARTIAL','Published v39: synthetic Grade 7 teacher count 3 saved; profile wheel displayed Awareness 3/6 Proficient. View mappings opened the exact approved BULK39 observation with Awareness and SCC1.1. Complete self/peer/teacher comparison, empty-versus-zero semantics and accessibility coverage remain open.'),
'Context, barrier and official score':('PARTIAL','Published v40: added scoped status controls. Synthetic milestone changed to completed and QA40 barrier changed from monitoring to resolved; both survived leaving and reopening HPC, with term/class retained. Cross-record/status negative route tests pass. Official mappings and final-score lifecycle remain incomplete.'),
'Contributor sharing lifecycle':('PARTIAL','Published v42 keeps a newly generated contribution URL visibly selectable if clipboard access fails. Route tests reject disabled schools and resolve the learner grade to the active stage framework instead of selecting an arbitrary framework. Existing link history renders. A new live named-contributor submission and browser revocation were not performed because they create and revoke externally usable access.'),
'Role and feature-flag enforcement':('PARTIAL','Handler tests deny Student/Parent/unknown roles, disabled HPC and foreign-school share creation; active teacher accepted. Public contribution now independently rechecks the school feature flag. Live multi-account/disabled-school tests remain.'),
'Parent/peer token isolation':('PARTIAL','Foreign-school share creation and foreign portfolio/certificate download are denied in handler tests. Handler tests reject tampered, expired and revoked tokens without writes; the contribution route now binds the framework to the learner grade. Live contributor identity and token-lifecycle checks remain unexecuted.'),
'Narrow layout and full navigation':('PASS','Published v45 mobile viewport measured 375px client and 375px document width after a real 36px overflow defect was repaired. The mobile shell, form fields and horizontal navigation remained usable in the sampled HPC state.'),
'Desktop/laptop/tablet/phone matrix':('PARTIAL','Published v45 measured phone 375px, tablet 753px and desktop 1425px with document width equal to client width at each size. Screenshots were inspected; tablet header clutter was repaired. The complete requested dimension/orientation set and physical devices remain outstanding.'),
'Keyboard, screen reader and real devices':('PARTIAL','Published v45 DOM audit found 121 form controls with accessible labels, 80 buttons with nonempty names and one page-level H1; sampled browser console contained no errors. Keyboard-only, 200% zoom, screen-reader and physical-device certification remain outstanding.'),
'Automated suite and build':('PASS','Published v45: 103 automated tests pass, including route security, stage-boundary, null-grade and responsive source regressions. Application TypeScript check and production build pass. Database/storage doubles do not establish live end-to-end acceptance.'),
}
cases=[(*c[:4],updates[c[1]][1],updates[c[1]][0]) if c[1] in updates else c for c in cases]
doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)
def style(name,size,before,after,color='202020',bold=False):
 s=doc.styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=bold;s.font.color.rgb=RGBColor.from_string(color)
 p=s.paragraph_format;p.space_before=Pt(before);p.space_after=Pt(after);p.line_spacing=1.1
 return s
style('Normal',11,0,6);style('Title',23,0,4,'202020',True);style('Subtitle',12,0,10,'555555')
style('Heading 1',16,16,8,'2E74B5',True);style('Heading 2',13,12,6,'2E74B5',True);style('Heading 3',12,8,4,'1F4D78',True)
for name in ['Heading 1','Heading 2','Heading 3']:doc.styles[name].paragraph_format.keep_with_next=True
header=sec.header.paragraphs[0];header.text='EDUAI  |  HPC acceptance testing';header.style='Caption'
header.runs[0].font.size=Pt(9)
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('1 September 2026  |  Page ').font.size=Pt(9)
field=OxmlElement('w:fldSimple');field.set(qn('w:instr'),'PAGE');footer._p.append(field)
doc.add_paragraph('HPC Test Cases & Results',style='Title')
doc.add_paragraph('Prompts 1–5 | Published version 45 + retest results | Acceptance register',style='Subtitle')
doc.add_paragraph('Decision: NOT READY FOR FULL SIGN-OFF',style='Heading 1')
doc.add_paragraph('This document records executed tests and the remaining acceptance plan. A passing scenario does not mean its entire prompt is complete. Google login is deferred at the user’s request. Critical stage/scoring and coverage gaps remain.')
counts=Counter(c[-1] for c in cases)
doc.add_paragraph(f'{len(cases)} acceptance cases: '+ '; '.join(f'{n} {s.lower()}' for s,n in sorted(counts.items()))+'. The separate automated suite has 103 passing checks; these counts must not be added together as one coverage percentage. Google remains deferred by request; blocked cases require external accounts, data or completed prerequisite evidence.')
doc.add_paragraph('Scope and evidence',style='Heading 1')
doc.add_paragraph('HPC site: https://eduai-learning-xray-hpc-test.accounts740459.chatgpt.site/app. Supabase: mncprowjqrtmuqvxrcqq. Published source: df8fb51db59b4011a2a00d5b7e2594434681beda. Versions 42-45 add contributor-stage safeguards, visible link fallback, null-grade protection and responsive layout repairs. The original live Learning X-Ray site/database were not modified.')
doc.add_paragraph('Evidence: direct browser actions and DOM read-back through 1 September 2026; phone, tablet and desktop screenshots; anonymous HTTP requests; automated test output; and earlier labelled execution notes. No fresh password login, physical-device or screen-reader certification, cross-school RLS certification, externally submitted contribution link, or successful final HPC report is claimed.')
doc.add_paragraph('Status definitions',style='Heading 1')
doc.add_paragraph('PASS: stated scenario met its expected result. PARTIAL: only part verified. FAIL: observed defect or confirmed unavailable behaviour. BLOCKED: prerequisite or environment prevents execution. NOT TESTED: no execution evidence yet. DEFERRED: explicitly paused by the user.')
doc.add_paragraph('Test data and cleanup',style='Heading 1')
doc.add_paragraph('Only synthetic HPC data was written: a Grade7 learner TEST ONLY QA Middle20260831 (roll QA-M-0831); academic goal; Science mapping; parent evidence awaiting review; excluded peer evidence; Awareness self-count3. Earlier same-day tests added an approved synthetic reflection and planned milestone to existing Grade10 test records. Records are retained for reproducibility; no cleanup deletion performed.')
doc.add_paragraph('Priority defects and blockers',style='Heading 1')
for text in [
 'Resolved in v37: Grade10 no longer uses the Middle-only six-statement form.',
 'High: complete stage-specific framework/outcome coverage and a successful finalization journey are unproven.',
 'Resolved in v37: existing-student linking and immediate learner refresh passed browser checks.',
 'Published and regression-tested: role and school-feature route gates, share ownership/stage selection, rubric hydration, annual draft generation and file/hour validation. Broader live multi-account acceptance remains.',
 'Resolved in v39: bulk observations now save per learner; Interventions displays HPC actions. Broader lifecycle tests remain partial. Microsoft provider configuration remains unavailable.',
 'Security sign-off blocked: two-school/role and live parent/peer identity isolation are not executed.',
 'Regression blockers: zero credits and no controlled fresh-login inbox account. Responsive viewport checks pass at sampled phone, tablet and desktop sizes; physical-device, keyboard and screen-reader certification remain.'
]:doc.add_paragraph(text)

mh=[('01','Versioned stage framework','FAIL','Stage mismatch remains; full boundary/version tests pending.'),('02','Learner context','FAIL','Creation works; existing-student linkage and full context validation incomplete.'),('03','Goals and reflection','NOT VERIFIED','Middle goal passes; all stage-specific sections untested.'),('04','Official mapping','NOT VERIFIED','One Middle Science mapping passes; complete outcomes/stages pending.'),('05','Activity and rubric','NOT VERIFIED','Complete creation/mapping lifecycle pending.'),('06','Teacher observation','NOT VERIFIED','Bulk and individual override/attachment flow pending.'),('07','Self-assessment','NOT VERIFIED','Reflection and one Middle count pass; all structures pending.'),('08','Peer feedback','NOT VERIFIED','Entry/exclusion sampled; assignment and identity/isolation pending.'),('09','Parent partnership','NOT VERIFIED','Teacher entry passes; external parent flow pending.'),('10','Portfolio and traceability','NOT VERIFIED','Upload/download and evidence-level links pending.'),('11','Official scoring/wheel','FAIL','Wrong-stage scoring observed; boundaries alone do not prove stage selection.'),('12','Strengths and support','NOT VERIFIED','Sensitive notes, override audit and intervention integration pending.'),('13','Applied learning','NOT VERIFIED','Detail/milestone pass; full project/course lifecycle pending.'),('14','Annual holistic profile','NOT VERIFIED','Full evidence-backed annual synthesis pending.'),('15','Finalization and reporting','NOT VERIFIED','Negative readiness passes; positive immutable finalization/PDF/role tests blocked.')]
mh=[(n,t,'NOT VERIFIED', 'Stage isolation and linking defects were fixed in v37; broader requirement acceptance remains incomplete.') if n in {'01','02','11'} else (n,t,s,w) for n,t,s,w in mh]
doc.add_paragraph('MH-01 to MH-15 acceptance matrix',style='Heading 1')
doc.add_paragraph('No must-have is signed off from partial scenarios. NOT VERIFIED is explicitly not a pass.')
for number,title,status,why in mh:
 doc.add_paragraph(f'MH-{number} | {title} | {status}',style='Heading 3');doc.add_paragraph(why)

area=None
for i,(group,title,steps,expected,actual,status) in enumerate(cases,1):
 if group!=area:doc.add_paragraph(group,style='Heading 1');area=group
 doc.add_paragraph(f'TC-{i:03d} | {title} | {status}',style='Heading 2')
 for label,text in [('Procedure',steps),('Expected',expected),('Actual',actual)]:
  p=doc.add_paragraph();p.add_run(label+': ').bold=True;p.add_run(text)
doc.add_paragraph('Retest and release gate',style='Heading 1')
doc.add_paragraph('Resolve high-severity defects first. Prepare two isolated school accounts, controlled email inboxes, principal and contributor test access, representative files and funded test credits. Retest all failed, blocked, partial and unexecuted cases on the published revision. Run real-device and supported browser checks. Sign off only when every mandatory requirement and security boundary is demonstrated; retain exact build, data and execution evidence.')
doc.add_paragraph('Source requirements',style='Heading 1')
doc.add_paragraph('Manager-provided HPC Integration Codex Prompts v2-1 (document heading: Verified v2.0,19August2026), Prompts1–5; HPC Must Have Specification v2-1, MH-01–MH-15. Repository test suite and same-day acceptance notes supply implementation evidence, not independent official framework approval.')
doc.core_properties.title='HPC Test Cases and Results — Prompts 1–5'
doc.core_properties.subject='Interim acceptance register; not final sign-off'
doc.core_properties.author='EduAI QA'
path=OUT/'HPC-Test-Cases-and-Results-2026-08-31-Updated.docx';doc.save(path)
assert len(cases)==len(set(f'TC-{i:03}' for i in range(1,len(cases)+1)))
assert all(c[-1] in {'PASS','PARTIAL','FAIL','BLOCKED','NOT TESTED','DEFERRED'} for c in cases)
print(path);print(dict(counts));print('Cases:',len(cases))
